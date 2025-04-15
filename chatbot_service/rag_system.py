import os
import pickle
import logging
import numpy as np
from pathlib import Path
import nltk
from nltk.tokenize import sent_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import json
import re
import csv
import tempfile
import subprocess

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Download NLTK data if needed
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

# File to store document paths
DOC_REGISTRY_FILE = "document_registry.pkl"

class DocumentParser:
    """Handles parsing of different document formats into plain text"""
    
    @staticmethod
    def parse_document(file_path):
        """Parse a document and extract text based on its extension"""
        # Get the file extension (lowercase)
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # Handle different file formats
        if ext == '.txt':
            return DocumentParser._parse_txt(file_path)
        elif ext == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return DocumentParser._parse_doc(file_path)
        elif ext in ['.csv', '.xls', '.xlsx']:
            return DocumentParser._parse_tabular(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    @staticmethod
    def _parse_txt(file_path):
        """Parse a plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encodings
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error reading text file: {e}")
                raise
    
    @staticmethod
    def _parse_pdf(file_path):
        """Parse a PDF file"""
        try:
            # Try to import PyPDF2
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except ImportError:
            logger.warning("PyPDF2 not installed, trying external tools...")
            return DocumentParser._parse_pdf_external(file_path)
        except Exception as e:
            logger.error(f"Error with PyPDF2: {e}")
            return DocumentParser._parse_pdf_external(file_path)
    
    @staticmethod
    def _parse_pdf_external(file_path):
        """Parse PDF using external tools if available"""
        try:
            # Try using pdftotext from poppler if available
            with tempfile.NamedTemporaryFile(suffix='.txt') as temp_txt:
                subprocess.run(['pdftotext', file_path, temp_txt.name], check=True)
                with open(temp_txt.name, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception as e:
            logger.error(f"Error using external PDF tool: {e}")
            raise ValueError("Could not parse PDF. Please install PyPDF2 or ensure pdftotext is in your PATH.")
    
    @staticmethod
    def _parse_doc(file_path):
        """Parse a DOC/DOCX file"""
        try:
            # Try to import docx
            import docx
            
            text = ""
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except ImportError:
            logger.error("python-docx not installed")
            raise ValueError("python-docx is not installed. Please install it with 'pip install python-docx'")
        except Exception as e:
            logger.error(f"Error parsing DOC/DOCX: {e}")
            raise
    
    @staticmethod
    def _parse_tabular(file_path):
        """Parse a tabular file (CSV, XLS, XLSX)"""
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        try:
            if ext == '.csv':
                return DocumentParser._parse_csv(file_path)
            else:
                return DocumentParser._parse_excel(file_path)
        except Exception as e:
            logger.error(f"Error parsing tabular file: {e}")
            raise
    
    @staticmethod
    def _parse_csv(file_path):
        """Parse a CSV file"""
        text = ""
        try:
            # First try to detect the delimiter
            with open(file_path, 'r', encoding='utf-8') as f:
                sample = f.read(4096)  # Read a sample to detect delimiter
                
                # Count potential delimiters in the sample
                delimiters = [',', ';', '\t', '|']
                counts = {d: sample.count(d) for d in delimiters}
                
                # Use the most common delimiter
                delimiter = max(counts.items(), key=lambda x: x[1])[0]
                if counts[delimiter] == 0:
                    delimiter = ','  # Default to comma if none found
            
            # Now parse the CSV with the detected delimiter
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f, delimiter=delimiter)
                headers = next(reader, None)  # Try to get headers
                
                if headers:
                    # Add headers as the first line
                    text += " | ".join(headers) + "\n"
                    
                    # Add each row
                    for row in reader:
                        if any(cell.strip() for cell in row):  # Skip empty rows
                            text += " | ".join(row) + "\n"
                
            return text
        except Exception as e:
            logger.error(f"Error parsing CSV: {e}")
            # Fallback to simple reading
            return DocumentParser._parse_txt(file_path)
    
    @staticmethod
    def _parse_excel(file_path):
        """Parse an Excel file"""
        try:
            # Try to import pandas
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            text = ""
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                # Format sheet header
                text += f"Sheet: {sheet_name}\n"
                
                # Convert dataframe to text representation
                text += df.to_string(index=False) + "\n\n"
            
            return text
        except ImportError:
            logger.error("pandas not installed")
            raise ValueError("pandas is not installed. Please install it with 'pip install pandas'")
        except Exception as e:
            logger.error(f"Error parsing Excel: {e}")
            raise

class RAGSystem:
    """Retrieval Augmented Generation system for document knowledge base"""
    
    def __init__(self, docs_folder="documents"):
        self.docs_folder = docs_folder
        self.documents = {}  # Store document content
        self.document_paths = []  # Store paths to registered documents
        self.vectorizer = TfidfVectorizer(stop_words='english')
        self.document_vectors = None
        self.chunk_map = []  # Maps index to (doc_name, chunk_index)
        
        # Create documents folder if it doesn't exist
        os.makedirs(self.docs_folder, exist_ok=True)
        
        # Load document registry if it exists
        self.load_document_registry()
        
        # Load existing documents
        self.load_documents()
    
    def load_document_registry(self):
        """Load the registry of document paths"""
        if os.path.exists(DOC_REGISTRY_FILE):
            try:
                with open(DOC_REGISTRY_FILE, 'rb') as f:
                    self.document_paths = pickle.load(f)
                logger.info(f"Loaded {len(self.document_paths)} document paths from registry")
            except Exception as e:
                logger.error(f"Error loading document registry: {e}")
                self.document_paths = []
        else:
            self.document_paths = []
    
    def save_document_registry(self):
        """Save the registry of document paths"""
        try:
            with open(DOC_REGISTRY_FILE, 'wb') as f:
                pickle.dump(self.document_paths, f)
            logger.info(f"Saved {len(self.document_paths)} document paths to registry")
        except Exception as e:
            logger.error(f"Error saving document registry: {e}")
    
    def load_documents(self):
        """Load all documents from the documents folder and registry"""
        # First load documents from the documents folder
        supported_extensions = ['.txt', '.pdf', '.doc', '.docx', '.csv', '.xls', '.xlsx']
        for ext in supported_extensions:
            for file_path in Path(self.docs_folder).glob(f"*{ext}"):
                self.add_document(str(file_path), register=False)  # Don't re-register folder docs
        
        # Then load documents from the registry
        for file_path in self.document_paths:
            if os.path.exists(file_path):
                # Skip if already loaded from documents folder
                folder_files = []
                for ext in supported_extensions:
                    folder_files.extend([str(p) for p in Path(self.docs_folder).glob(f"*{ext}")])
                
                if file_path not in folder_files:
                    self.add_document(file_path, register=False)  # Already registered
        
        # Vectorize all documents
        if self.documents:
            self._vectorize_documents()
            logger.info(f"Loaded {len(self.documents)} documents into RAG system")
    
    def add_document(self, file_path, register=True):
        """Add a document to the knowledge base"""
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"Error: File not found at '{file_path}'")
            return False
            
        try:
            # Print absolute path to help debug
            abs_path = os.path.abspath(file_path)
            logger.info(f"Loading document from: {abs_path}")
            
            # Get file extension
            _, ext = os.path.splitext(abs_path)
            ext = ext.lower()
            
            # Check supported extensions
            supported_extensions = ['.txt', '.pdf', '.doc', '.docx', '.csv', '.xls', '.xlsx']
            if ext not in supported_extensions:
                logger.error(f"Error: Unsupported file format {ext}")
                return False
            
            # Parse document based on type
            try:
                text = DocumentParser.parse_document(abs_path)
            except ValueError as e:
                logger.error(f"Error parsing document: {e}")
                return False
            except Exception as e:
                logger.error(f"Unexpected error parsing document: {e}")
                return False
            
            # Store document by filename
            filename = os.path.basename(file_path)
            self.documents[filename] = self._split_into_chunks(text)
            
            # Add to document registry if not already there
            if register and abs_path not in self.document_paths:
                self.document_paths.append(abs_path)
                self.save_document_registry()

            self._vectorize_documents()
            
            logger.info(f"Added document: {filename}")
            logger.info(f"Extracted {len(text)} characters, split into {len(self.documents[filename])} chunks")
            return True
        except Exception as e:
            logger.error(f"Error adding document {file_path}: {str(e)}")
            return False
    
    def _split_into_chunks(self, text, max_chunk_size=1000):
        """Split text into manageable chunks based on document type"""
        # Check if text has tabular format (indication of CSV/Excel)
        is_tabular = "|" in text and "\n" in text and text.count("|") > text.count("\n")
        
        if is_tabular:
            # For tabular data, split by sheets/sections and maintain row structure
            chunks = []
            current_section = ""
            section_header = None
            rows = text.split("\n")
            
            for row in rows:
                # Check for section headers (like "Sheet: SheetName")
                if row.strip().startswith("Sheet:"):
                    # If we were building a section, add it to chunks
                    if current_section:
                        chunks.append(current_section)
                    # Start new section
                    section_header = row
                    current_section = row + "\n"
                    continue
                    
                # If this row would make the section too large, create a new chunk
                # But keep the section header for context
                if len(current_section) + len(row) > max_chunk_size and current_section:
                    chunks.append(current_section)
                    current_section = (section_header + "\n") if section_header else ""
                
                # Add the row to the current section
                if row.strip():  # Skip empty rows
                    current_section += row + "\n"
            
            # Add the last section if it has content
            if current_section:
                chunks.append(current_section)
        else:
            # For regular text documents, use sentence-based chunking
            sentences = sent_tokenize(text)
            chunks = []
            current_chunk = ""
            
            for sentence in sentences:
                # If adding this sentence would exceed max size, store current chunk and start new one
                if len(current_chunk) + len(sentence) > max_chunk_size and current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = sentence
                else:
                    current_chunk += " " + sentence if current_chunk else sentence
            
            # Add the last chunk if it has content
            if current_chunk:
                chunks.append(current_chunk)
                
        return chunks
        
    def _vectorize_documents(self):
        """Create vector representations of all document chunks"""
        all_chunks = []
        self.chunk_map = []  # Maps index to (doc_name, chunk_index)
        
        for doc_name, chunks in self.documents.items():
            for i, chunk in enumerate(chunks):
                all_chunks.append(chunk)
                self.chunk_map.append((doc_name, i))
        
        if all_chunks:
            self.document_vectors = self.vectorizer.fit_transform(all_chunks)
            logger.info(f"Vectorized {len(all_chunks)} document chunks")
        else:
            self.document_vectors = None
            logger.warning("No document chunks to vectorize")
    
    def query(self, query_text, top_k=3):
        """Find the most relevant document chunks for a query"""
        if not self.documents or self.document_vectors is None:
            logger.warning("No documents available for querying")
            return []
        
        # Expand query with related terms for better retrieval
        expanded_query = query_text
        
        # Add restaurant-specific terms to improve retrieval
        if any(keyword in query_text.lower() for keyword in ['recommend', 'suggestion', 'best']):
            expanded_query += " menu popular dishes signature specialties"
        elif 'wine' in query_text.lower() or 'drink' in query_text.lower():
            expanded_query += " beverages wines cocktails"
        elif 'allerg' in query_text.lower() or 'diet' in query_text.lower():
            expanded_query += " allergens dietary vegetarian vegan gluten"
            
        # Vectorize the query
        query_vector = self.vectorizer.transform([expanded_query])
        
        # Calculate similarity with all chunks
        similarities = cosine_similarity(query_vector, self.document_vectors).flatten()
        
        # Get indices of top_k most similar chunks
        top_indices = similarities.argsort()[-top_k:][::-1]
        
        # Return the relevant chunks with their similarity scores
        results = []
        for idx in top_indices:
            doc_name, chunk_idx = self.chunk_map[idx]
            chunk_text = self.documents[doc_name][chunk_idx]
            similarity = similarities[idx]
            
            # Only include chunks with meaningful similarity
            if similarity > 0.05:  # Lower threshold to ensure we get some results
                results.append({
                    "document": doc_name,
                    "text": chunk_text,
                    "similarity": round(float(similarity), 3)
                })
        
        return results
    
    def get_relevant_context(self, query, max_tokens=2000):
        """Get combined relevant context for a query within token limit"""
        # First try to get exact matches
        results = self.query(query, top_k=5)
        
        # Special handling for different query types
        is_tabular_query = any(word in query.lower() for word in ['excel', 'csv', 'table', 'row', 'column', 'sheet'])
        is_menu_query = any(word in query.lower() for word in ['menu', 'dish', 'food', 'recommend', 'price'])
        
        # For table-specific queries, prioritize tabular documents
        if is_tabular_query:
            # Boost relevance of table-formatted chunks
            for result in results:
                # Check if the chunk has table-like structure (multiple | characters)
                if result["text"].count("|") > 3:
                    result["similarity"] += 0.2  # Boost similarity score
            
            # Re-sort results based on adjusted similarity
            results = sorted(results, key=lambda x: x["similarity"], reverse=True)[:5]
        
        # Special handling for menu queries
        if is_menu_query:
            menu_chunks = []
            for doc_name, chunks in self.documents.items():
                if any(word in doc_name.lower() for word in ['menu', 'dish', 'food', 'price']):
                    for i, chunk in enumerate(chunks):
                        menu_chunks.append({
                            "document": doc_name,
                            "text": chunk,
                            "similarity": 0.9  # High but not maximum to allow for better exact matches
                        })
            
            # Add menu chunks first, then other results
            all_results = menu_chunks + [r for r in results if not any(word in r["document"].lower() 
                                                                    for word in ['menu', 'dish', 'food'])]
            results = sorted(all_results, key=lambda x: x["similarity"], reverse=True)[:5]
        
        context = ""
        
        for result in results:
            # Simple approximation: 1 token ≈ 4 characters
            estimated_tokens = len(context) / 4
            result_tokens = len(result["text"]) / 4
            
            if estimated_tokens + result_tokens < max_tokens:
                # For tabular data, format it nicely
                is_tabular = result["text"].count("|") > 3
                
                if is_tabular:
                    context += f"\nFrom {result['document']} (relevance: {result['similarity']}):\n```\n{result['text'].strip()}\n```\n"
                else:
                    context += f"\nFrom {result['document']} (relevance: {result['similarity']}):\n{result['text'].strip()}\n"
            else:
                break
                
        return context.strip()